"""Convert USER_GUIDE.md to USER_GUIDE.docx.

Lightweight Markdown -> docx converter tailored for this guide.
Supports: H1-H4, paragraphs, bullet lists, numbered lists, tables,
fenced code blocks, blockquotes, inline **bold**, *italic*, `code`.
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "USER_GUIDE.md"
DST = ROOT / "USER_GUIDE.docx"

INLINE_RE = re.compile(
    r"(\*\*(?P<bold>[^*]+)\*\*|\*(?P<italic>[^*]+)\*|`(?P<code>[^`]+)`|\[(?P<ltext>[^\]]+)\]\((?P<lhref>[^)]+)\))"
)


def add_inline(paragraph, text: str):
    """Add a text string with inline markdown styling to a paragraph."""
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos : m.start()])
        if m.group("bold"):
            r = paragraph.add_run(m.group("bold"))
            r.bold = True
        elif m.group("italic"):
            r = paragraph.add_run(m.group("italic"))
            r.italic = True
        elif m.group("code"):
            r = paragraph.add_run(m.group("code"))
            r.font.name = "Consolas"
            r.font.size = Pt(10)
        elif m.group("ltext"):
            # render link text only (plain word doc, no hyperlink wiring needed for a TOC-style list)
            r = paragraph.add_run(m.group("ltext"))
            r.font.color.rgb = RGBColor(0x0B, 0x5C, 0xAB)
            r.underline = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def is_table_separator(line: str) -> bool:
    s = line.strip().strip("|")
    if not s:
        return False
    cells = [c.strip() for c in s.split("|")]
    return all(re.fullmatch(r":?-{2,}:?", c) for c in cells) and len(cells) >= 1


def split_table_row(line: str) -> list[str]:
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip() for c in s.split("|")]


def build():
    md = SRC.read_text(encoding="utf-8").splitlines()
    doc = Document()

    # Base style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    i = 0
    n = len(md)
    while i < n:
        line = md[i]
        stripped = line.strip()

        # Blank line
        if not stripped:
            i += 1
            continue

        # Horizontal rule -> skip (acts as section divider)
        if re.fullmatch(r"-{3,}", stripped):
            i += 1
            continue

        # Headings
        h = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if h:
            level = len(h.group(1))
            text = h.group(2).strip()
            # Strip trailing markdown link markup inside heading if any
            heading = doc.add_heading(level=min(level, 4))
            add_inline(heading, text)
            i += 1
            continue

        # Fenced code block
        if stripped.startswith("```"):
            i += 1
            code_lines = []
            while i < n and not md[i].strip().startswith("```"):
                code_lines.append(md[i])
                i += 1
            i += 1  # skip closing fence
            p = doc.add_paragraph()
            r = p.add_run("\n".join(code_lines))
            r.font.name = "Consolas"
            r.font.size = Pt(10)
            continue

        # Blockquote
        if stripped.startswith(">"):
            text = stripped.lstrip(">").strip()
            p = doc.add_paragraph(style="Intense Quote") if "Intense Quote" in [s.name for s in doc.styles] else doc.add_paragraph()
            add_inline(p, text)
            i += 1
            continue

        # Table
        if stripped.startswith("|") and i + 1 < n and is_table_separator(md[i + 1]):
            header = split_table_row(md[i])
            i += 2  # skip header + separator
            rows = []
            while i < n and md[i].strip().startswith("|"):
                rows.append(split_table_row(md[i]))
                i += 1
            table = doc.add_table(rows=1 + len(rows), cols=len(header))
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            for c, text in enumerate(header):
                p = hdr[c].paragraphs[0]
                add_inline(p, text)
                for run in p.runs:
                    run.bold = True
            for ri, row in enumerate(rows, start=1):
                cells = table.rows[ri].cells
                for c, text in enumerate(row):
                    if c >= len(cells):
                        break
                    para = cells[c].paragraphs[0]
                    add_inline(para, text)
            doc.add_paragraph()
            continue

        # Unordered list
        if re.match(r"^[-*]\s+", stripped):
            while i < n and re.match(r"^[-*]\s+", md[i].strip()):
                text = re.sub(r"^[-*]\s+", "", md[i].strip())
                p = doc.add_paragraph(style="List Bullet")
                add_inline(p, text)
                i += 1
            continue

        # Ordered list
        if re.match(r"^\d+\.\s+", stripped):
            while i < n and re.match(r"^\d+\.\s+", md[i].strip()):
                text = re.sub(r"^\d+\.\s+", "", md[i].strip())
                p = doc.add_paragraph(style="List Number")
                add_inline(p, text)
                i += 1
            continue

        # Paragraph (collect until blank line)
        para_lines = [stripped]
        i += 1
        while i < n and md[i].strip() and not re.match(
            r"^(#{1,6}\s|[-*]\s|\d+\.\s|\||>|```|-{3,}$)", md[i].strip()
        ):
            para_lines.append(md[i].strip())
            i += 1
        p = doc.add_paragraph()
        add_inline(p, " ".join(para_lines))

    doc.save(DST)
    print(f"Wrote {DST}")


if __name__ == "__main__":
    build()
